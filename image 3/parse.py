from docx import Document

def extract_text_from_docx(file_path):
    """Extracts both paragraph text and table data from a Word document."""
    doc = Document(file_path)
    
    # Extract paragraph text
    extracted_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            extracted_text.append(text)
    
    # Extract table data
    tables_data = []
    for table in doc.tables:
        table_content = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]  # Extract text from each cell
            table_content.append(row_data)
        tables_data.append(table_content)
    
    return extracted_text, tables_data

def format_timetable(table_data):
    """Formats extracted timetable table into a readable dictionary."""
    timetable = {}
    headers = table_data[0]  # First row contains headers like "Day", "8:30-9:25", etc.
    
    for row in table_data[1:]:  # Skip header row
        if row:  
            day = row[0]  # First column contains the day (e.g., Monday, Tuesday)
            classes = row[1:]  # Remaining columns contain class details
            timetable[day] = {headers[i + 1]: classes[i] for i in range(len(classes))}
    
    return timetable

# Example Usage
file_path = r"C:\Users\Ravi\Desktop\test2.doc.docx" # Change this to the actual document path
text_data, tables_data = extract_text_from_docx(file_path)

# Print extracted text (paragraphs)
print("\n--- Extracted Paragraph Text ---\n")
for line in text_data:
    print(line)

# Print extracted tables
print("\n--- Extracted Table Data (Raw) ---\n")
for table in tables_data:
    for row in table:
        print(row)

# If the document contains a timetable, format and print it
if tables_data:
    formatted_timetable = format_timetable(tables_data[0])  # Assume first table is the timetable
    print("\n--- Formatted Timetable ---\n")
    for day, schedule in formatted_timetable.items():
        print(f"{day}:")
        for time, subject in schedule.items():
            print(f"  {time}: {subject}")
